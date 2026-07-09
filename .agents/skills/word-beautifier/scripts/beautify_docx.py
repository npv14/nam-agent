import docx
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.text.run import Run
import re
import sys

# Helper functions for raw XML formatting extraction

def is_xml_run_bold(r_elem):
    rPr = [child for child in r_elem if child.tag.endswith('rPr')]
    if rPr:
        b = [child for child in rPr[0] if child.tag.endswith('b')]
        if b:
            val = b[0].get(qn('w:val'))
            if val in ('false', '0', 'none', 'off'):
                return False
            return True
    return False

def is_xml_run_italic(r_elem):
    rPr = [child for child in r_elem if child.tag.endswith('rPr')]
    if rPr:
        i = [child for child in rPr[0] if child.tag.endswith('i')]
        if i:
            val = i[0].get(qn('w:val'))
            if val in ('false', '0', 'none', 'off'):
                return False
            return True
    return False

def is_xml_run_underline(r_elem):
    rPr = [child for child in r_elem if child.tag.endswith('rPr')]
    if rPr:
        u = [child for child in rPr[0] if child.tag.endswith('u')]
        if u:
            val = u[0].get(qn('w:val'))
            if val in ('false', '0', 'none', 'off'):
                return False
            return True
    return False

def get_xml_run_text(r_elem):
    t_elms = [child for child in r_elem if child.tag.endswith('t')]
    return "".join(t.text for t in t_elms if t.text)

# Helper functions for styling

def set_cell_background(cell, hex_color):
    shd_elms = [child for child in cell._tc.get_or_add_tcPr() if child.tag.endswith('shd')]
    for shd in shd_elms:
        cell._tc.tcPr.remove(shd)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar_elms = [child for child in tcPr if child.tag.endswith('tcMar')]
    for tm in tcMar_elms:
        tcPr.remove(tm)
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders_elms = [child for child in tcPr if child.tag.endswith('tcBorders')]
    for tb in tcBorders_elms:
        tcPr.remove(tb)
        
    tcBorders = OxmlElement('w:tcBorders')
    for border_name, border_style in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if border_style:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), border_style.get('val', 'single'))
            border.set(qn('w:sz'), str(border_style.get('sz', 4)))
            border.set(qn('w:space'), str(border_style.get('space', 0)))
            border.set(qn('w:color'), border_style.get('color', 'auto'))
            tcBorders.append(border)
        else:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tcBorders.append(border)
    tcPr.append(tcBorders)

def set_paragraph_left_border(paragraph, color_hex="E1B924", size=36, space=8):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr_elms = [child for child in pPr if child.tag.endswith('pBdr')]
    if pBdr_elms:
        pBdr = pBdr_elms[0]
    else:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    
    left_elms = [child for child in pBdr if child.tag.endswith('left')]
    for left in left_elms:
        pBdr.remove(left)
        
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(size))
    left.set(qn('w:space'), str(space))
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)

def add_bookmark(paragraph, bookmark_id, bookmark_name):
    p = paragraph._element
    pPr = p.pPr
    insert_idx = p.index(pPr) + 1 if pPr is not None else 0
    
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(bookmark_id))
    bookmark_start.set(qn('w:name'), bookmark_name)
    
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(bookmark_id))
    
    p.insert(insert_idx, bookmark_start)
    p.append(bookmark_end)

def add_hyperlink(paragraph, url, text, font_name="Segoe UI", font_size=11, color_rgb=(0, 48, 135), bold=False, underline=True):
    part = paragraph.part
    rId = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    p = paragraph._element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), rId)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size * 2)))
    rPr.append(sz)
    
    color = OxmlElement('w:color')
    color_hex = f"{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}"
    color.set(qn('w:val'), color_hex)
    rPr.append(color)
    
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
        
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
    new_run.append(rPr)
    
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    p.append(hyperlink)

def add_hyperlink_anchor(paragraph, anchor_name, text, font_name="Segoe UI", font_size=11, color_rgb=(0, 48, 135), bold=False, underline=True):
    p = paragraph._element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), anchor_name)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size * 2)))
    rPr.append(sz)
    
    color = OxmlElement('w:color')
    color_hex = f"{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}"
    color.set(qn('w:val'), color_hex)
    rPr.append(color)
    
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
        
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
    new_run.append(rPr)
    
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    p.append(hyperlink)

def copy_paragraph_runs(src_p, dest_p, font_name="Segoe UI", font_size=11, color_rgb=(43, 43, 43)):
    for child in src_p._element:
        if child.tag.endswith('r'):
            text = get_xml_run_text(child)
            if not text:
                continue
            new_run = dest_p.add_run(text)
            new_run.font.name = font_name
            new_run.font.size = Pt(font_size)
            new_run.font.color.rgb = RGBColor(*color_rgb)
            new_run.bold = is_xml_run_bold(child)
            new_run.italic = is_xml_run_italic(child)
            new_run.underline = is_xml_run_underline(child)
        elif child.tag.endswith('hyperlink'):
            rId = child.get(qn('r:id'))
            anchor = child.get(qn('w:anchor'))
            
            hl_runs_text = []
            bold = False
            italic = False
            for sub in child:
                if sub.tag.endswith('r'):
                    text = get_xml_run_text(sub)
                    hl_runs_text.append(text)
                    if is_xml_run_bold(sub):
                         bold = True
                    if is_xml_run_italic(sub):
                         italic = True
            hl_text = "".join(hl_runs_text)
            if not hl_text.strip():
                continue
                
            if rId:
                try:
                    rel = src_p.part.rels[rId]
                    url = rel.target_ref
                    add_hyperlink(dest_p, url, hl_text, font_name=font_name, font_size=font_size, color_rgb=(0, 48, 135), bold=bold, underline=True)
                except Exception:
                    new_run = dest_p.add_run(hl_text)
                    new_run.font.name = font_name
                    new_run.font.size = Pt(font_size)
                    new_run.font.color.rgb = RGBColor(*color_rgb)
                    new_run.bold = bold
                    new_run.italic = italic
            elif anchor:
                add_hyperlink_anchor(dest_p, anchor, hl_text, font_name=font_name, font_size=font_size, color_rgb=(0, 48, 135), bold=bold, underline=True)

def is_list_paragraph(p):
    style_name = p.style.name if p.style else 'Normal'
    if 'List' in style_name or 'Bullet' in style_name:
        return True
    
    pPr_elms = [child for child in p._element if child.tag.endswith('pPr')]
    if pPr_elms:
        numPr_elms = [child for child in pPr_elms[0] if child.tag.endswith('numPr')]
        if numPr_elms:
            return True
    return False

# Metric cards layout helper
def style_metric_cell(cell, value_text, label_text):
    set_cell_background(cell, "F4F7FB")
    set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
    border_left = {'val': 'single', 'sz': 24, 'space': 0, 'color': 'E1B924'}
    set_cell_borders(cell, left=border_left)
    
    p_val = cell.paragraphs[0]
    p_val.paragraph_format.space_before = Pt(0)
    p_val.paragraph_format.space_after = Pt(2)
    p_val.paragraph_format.line_spacing = 1.0
    run_val = p_val.add_run(value_text)
    run_val.font.name = 'Segoe UI'
    run_val.font.size = Pt(22)
    run_val.font.bold = True
    run_val.font.color.rgb = RGBColor(0, 48, 135)
    
    p_lbl = cell.add_paragraph()
    p_lbl.paragraph_format.space_before = Pt(0)
    p_lbl.paragraph_format.space_after = Pt(0)
    p_lbl.paragraph_format.line_spacing = 1.1
    run_lbl = p_lbl.add_run(label_text)
    run_lbl.font.name = 'Segoe UI'
    run_lbl.font.size = Pt(9.0)
    run_lbl.font.color.rgb = RGBColor(60, 60, 60)

# Icon cards layout helper
def style_icon_cell(cell, emoji, label):
    set_cell_background(cell, "F4F7FB")
    set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
    border_style = {'val': 'single', 'sz': 4, 'space': 0, 'color': 'E0E0E0'}
    set_cell_borders(cell, top=border_style, bottom=border_style, left=border_style, right=border_style)
    
    p_emoji = cell.paragraphs[0]
    p_emoji.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p_emoji.paragraph_format.space_before = Pt(0)
    p_emoji.paragraph_format.space_after = Pt(2)
    run_em = p_emoji.add_run(emoji)
    run_em.font.name = 'Segoe UI Emoji'
    run_em.font.size = Pt(20)
    
    p_lbl = cell.add_paragraph()
    p_lbl.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p_lbl.paragraph_format.space_before = Pt(0)
    p_lbl.paragraph_format.space_after = Pt(0)
    run_lbl = p_lbl.add_run(label)
    run_lbl.font.name = 'Segoe UI'
    run_lbl.font.size = Pt(9.5)
    run_lbl.font.bold = True
    run_lbl.font.color.rgb = RGBColor(43, 43, 43)

def set_table_widths(table, widths):
    for i, col in enumerate(table.columns):
        if i < len(widths):
            col.width = widths[i]
            for cell in col.cells:
                cell.width = widths[i]

def get_table_alignments(table_index, num_cols):
    align_map = {
        0: ['L', 'C', 'C', 'C'], 
        1: ['L', 'L', 'R', 'L'], 
        2: ['L', 'L', 'C'],      
        3: ['L', 'C', 'L', 'R'], 
        4: ['L', 'C', 'L', 'R'], 
        5: ['L', 'L', 'C', 'C', 'C', 'R'], 
        6: ['L', 'C', 'R', 'L'], 
        7: ['L', 'R', 'R', 'R'], 
        8: ['L', 'R', 'R', 'L'], 
        9: ['L', 'L', 'L', 'C', 'L'], 
        10: ['L', 'R', 'R'],    
        11: ['L', 'R'],         
        12: ['C', 'L', 'L']     
    }
    return align_map.get(table_index, ['L'] * num_cols)

def get_table_widths(table_index):
    width_map = {
        0: [Inches(2.5), Inches(1.2), Inches(1.2), Inches(1.37)],
        1: [Inches(2.2), Inches(1.2), Inches(1.2), Inches(1.67)],
        2: [Inches(2.0), Inches(3.27), Inches(1.0)],
        3: [Inches(2.2), Inches(1.0), Inches(1.8), Inches(1.27)],
        4: [Inches(2.2), Inches(1.0), Inches(1.8), Inches(1.27)],
        5: [Inches(1.8), Inches(1.8), Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.87)],
        6: [Inches(2.0), Inches(1.2), Inches(1.5), Inches(1.57)],
        7: [Inches(1.8), Inches(1.4), Inches(1.4), Inches(1.67)],
        8: [Inches(2.0), Inches(1.5), Inches(1.5), Inches(1.27)],
        9: [Inches(1.2), Inches(1.5), Inches(1.5), Inches(0.87), Inches(1.2)],
        10: [Inches(1.5), Inches(2.37), Inches(2.4)],
        11: [Inches(3.0), Inches(3.27)],
        12: [Inches(0.5), Inches(2.5), Inches(3.27)]
    }
    return width_map.get(table_index, None)

def style_table(table, col_alignments=None):
    tblPr = table._tbl.tblPr
    tblBorders_elms = [child for child in tblPr if child.tag.endswith('tblBorders')]
    for tb in tblBorders_elms:
        tblPr.remove(tb)
        
    for r_idx, row in enumerate(table.rows):
        is_header = (r_idx == 0)
        
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        
        if is_header:
            trPr.append(OxmlElement('w:tblHeader'))
            
        for c_idx, cell in enumerate(row.cells):
            top_pad = 140 if is_header else 110
            bot_pad = 140 if is_header else 110
            set_cell_margins(cell, top=top_pad, bottom=bot_pad, left=150, right=150)
            
            if is_header:
                set_cell_background(cell, "003087")
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F4F7FB")
                else:
                    set_cell_background(cell, "FFFFFF")
            
            border_top = {'val': 'single', 'sz': 4, 'space': 0, 'color': 'D0D0D0'} if not is_header else {'val': 'single', 'sz': 8, 'space': 0, 'color': '003087'}
            border_bottom = {'val': 'single', 'sz': 8, 'space': 0, 'color': '003087'} if is_header else {'val': 'single', 'sz': 4, 'space': 0, 'color': 'E0E0E0'}
            if r_idx == len(table.rows) - 1:
                border_bottom = {'val': 'single', 'sz': 8, 'space': 0, 'color': '003087'}
                
            set_cell_borders(cell, top=border_top, bottom=border_bottom)
            
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                
                alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                if col_alignments and c_idx < len(col_alignments):
                    align_char = col_alignments[c_idx]
                    if align_char == 'C':
                        alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    elif align_char == 'R':
                        alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                p.alignment = alignment
                
                for run in p.runs:
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(9.5) if not is_header else Pt(10)
                    if is_header:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        run.font.color.rgb = RGBColor(43, 43, 43)

def add_page_number_fields(paragraph):
    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p = paragraph._element
    
    r_trang = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Segoe UI')
    rFonts.set(qn('w:hAnsi'), 'Segoe UI')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '18')
    rPr.append(sz)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '666666')
    rPr.append(color)
    
    r_trang.append(rPr)
    t1 = OxmlElement('w:t')
    t1.text = "Trang "
    t1.set(qn('xml:space'), 'preserve')
    r_trang.append(t1)
    p.append(r_trang)
    
    fldSimple1 = OxmlElement('w:fldSimple')
    fldSimple1.set(qn('w:instr'), 'PAGE')
    r_fallback1 = OxmlElement('w:r')
    r_fallback1.append(rPr)
    t_fallback1 = OxmlElement('w:t')
    t_fallback1.text = "1"
    r_fallback1.append(t_fallback1)
    fldSimple1.append(r_fallback1)
    p.append(fldSimple1)
    
    r_slash = OxmlElement('w:r')
    r_slash.append(rPr)
    t2 = OxmlElement('w:t')
    t2.text = " / "
    t2.set(qn('xml:space'), 'preserve')
    r_slash.append(t2)
    p.append(r_slash)
    
    fldSimple2 = OxmlElement('w:fldSimple')
    fldSimple2.set(qn('w:instr'), 'NUMPAGES')
    r_fallback2 = OxmlElement('w:r')
    r_fallback2.append(rPr)
    t_fallback2 = OxmlElement('w:t')
    t_fallback2.text = "1"
    r_fallback2.append(t_fallback2)
    fldSimple2.append(r_fallback2)
    p.append(fldSimple2)

def setup_header_footer(doc):
    section = doc.sections[0]
    
    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(4)
    hrun = hp.add_run("UWA × Việt Nam: Chương Trình Liên Kết & Cơ Hội Chuyển Tiếp")
    hrun.font.name = 'Segoe UI'
    hrun.font.size = Pt(8.5)
    hrun.font.italic = True
    hrun.font.color.rgb = RGBColor(120, 120, 120)
    
    hpPr = hp._element.get_or_add_pPr()
    hpBdr_elms = [child for child in hpPr if child.tag.endswith('pBdr')]
    if hpBdr_elms:
        hpBdr = hpBdr_elms[0]
    else:
        hpBdr = OxmlElement('w:pBdr')
        hpPr.append(hpBdr)
        
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '003087')
    hpBdr.append(bottom)
    
    # Footer
    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number_fields(fp)

def beautify_document(src_path, dest_path):
    src_doc = docx.Document(src_path)
    dest_doc = docx.Document()
    
    for s in dest_doc.sections:
        s.page_width = Inches(8.27)
        s.page_height = Inches(11.69)
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        
    setup_header_footer(dest_doc)
    
    # Title
    p_title = dest_doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(8)
    p_title.paragraph_format.keep_with_next = True
    run_title = p_title.add_run("🇻🇳 UWA × Việt Nam – Chương Trình Liên Kết & Cơ Hội Chuyển Tiếp")
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0, 48, 135)
    set_paragraph_left_border(p_title, color_hex="E1B924", size=36, space=10)
    
    p_sub = dest_doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(4)
    p_sub.paragraph_format.keep_with_next = True
    run_sub = p_sub.add_run("The University of Western Australia (QS #77 World 2026 · CRICOS 00126G)")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(11)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    p_date = dest_doc.add_paragraph()
    p_date.paragraph_format.space_before = Pt(0)
    p_date.paragraph_format.space_after = Pt(24)
    p_date.paragraph_format.keep_with_next = True
    run_date = p_date.add_run("📅 Ngày tạo báo cáo: 31/03/2026")
    run_date.font.name = 'Segoe UI'
    run_date.font.size = Pt(10)
    run_date.font.color.rgb = RGBColor(120, 120, 120)
    
    # Table of Contents
    p_toc_head = dest_doc.add_paragraph()
    p_toc_head.paragraph_format.space_before = Pt(12)
    p_toc_head.paragraph_format.space_after = Pt(8)
    p_toc_head.paragraph_format.keep_with_next = True
    run_toc_head = p_toc_head.add_run("📋 Mục lục")
    run_toc_head.font.name = 'Segoe UI'
    run_toc_head.font.size = Pt(14)
    run_toc_head.font.bold = True
    run_toc_head.font.color.rgb = RGBColor(0, 48, 135)
    
    toc_items = [
        ("sec_1", "1. Tổng quan quan hệ đối tác UWA – Việt Nam"),
        ("sec_2", "2. Chương trình 3+2 Dual Degree: UWA × VinUniversity"),
        ("sec_3", "3. Hợp tác Y tế: UWA × VinMec / VinUni Health Sciences"),
        ("sec_4", "4. Chương trình Pathways – UWA College (INTO Partnership)"),
        ("sec_5", "5. Assured Pathways – Chuyển tiếp nội bộ UWA (UG → PG)"),
        ("sec_6", "6. PhD & Nghiên cứu – Cơ hội cho sinh viên Việt Nam"),
        ("sec_7", "7. Western Australia Global Pathways (tại Việt Nam)"),
        ("sec_8", "8. Nguồn tham khảo")
    ]
    
    for anchor, text in toc_items:
        p_item = dest_doc.add_paragraph()
        p_item.paragraph_format.space_before = Pt(2)
        p_item.paragraph_format.space_after = Pt(3)
        p_item.paragraph_format.left_indent = Pt(15)
        
        run_bullet = p_item.add_run("• ")
        run_bullet.font.name = 'Segoe UI'
        run_bullet.font.size = Pt(11)
        run_bullet.font.color.rgb = RGBColor(0, 48, 135)
        
        add_hyperlink_anchor(p_item, anchor, text, font_name="Segoe UI", font_size=11, color_rgb=(0, 48, 135), bold=False, underline=True)
        
    dest_doc.add_page_break()
    
    # Loop body elements
    body = src_doc.element.body
    
    started = False
    skip_mode_stats = False
    skip_mode_icons = False
    table_index = 0
    bookmark_counter = 1
    
    active_heading_level = None
    
    for child in body:
        if child.tag.endswith('p'):
            p = Paragraph(child, src_doc)
            text = p.text.strip()
            
            if "1. Tổng quan quan hệ đối tác UWA" in text:
                started = True
                
            if not started:
                continue
                
            if skip_mode_stats:
                if "Tóm tắt quan hệ" in text:
                    skip_mode_stats = False
                else:
                    continue
                    
            if skip_mode_icons:
                if "5. Assured Pathways" in text:
                    skip_mode_icons = False
                else:
                    continue
                    
            style_name = p.style.name if p.style else 'Normal'
            
            if not text:
                if style_name.startswith('Heading'):
                    try:
                        active_heading_level = int(style_name.replace('Heading', '').strip())
                    except ValueError:
                        pass
                continue
                
            # Stats Card
            if text == "250+":
                stats = [
                    ("250+", "Đối tác quốc tế trên toàn cầu"),
                    ("45+", "Quốc gia có đối tác"),
                    ("100+", "Bài báo đồng tác giả UWA–VN (5 năm qua)"),
                    ("2024", "Năm ký MOA UWA × VinUni")
                ]
                metric_table = dest_doc.add_table(rows=1, cols=4)
                set_table_widths(metric_table, [Inches(1.5), Inches(1.5), Inches(1.6), Inches(1.67)])
                for s_idx, (val, lbl) in enumerate(stats):
                    cell = metric_table.cell(0, s_idx)
                    style_metric_cell(cell, val, lbl)
                p_space = dest_doc.add_paragraph()
                p_space.paragraph_format.space_before = Pt(0)
                p_space.paragraph_format.space_after = Pt(12)
                skip_mode_stats = True
                continue
                
            # Icon Card
            if text == "🏗️" or text.startswith("🏗️"):
                icons = [
                    ("🏗️", "Kiến trúc"),
                    ("💼", "Kinh doanh"),
                    ("⚙️", "Kỹ thuật"),
                    ("💻", "Khoa học Máy tính"),
                    ("🔬", "Khoa học Đời sống / Vật lý")
                ]
                icon_table = dest_doc.add_table(rows=1, cols=5)
                set_table_widths(icon_table, [Inches(1.25)] * 5)
                for i_idx, (em, lbl) in enumerate(icons):
                    cell = icon_table.cell(0, i_idx)
                    style_icon_cell(cell, em, lbl)
                p_space = dest_doc.add_paragraph()
                p_space.paragraph_format.space_before = Pt(0)
                p_space.paragraph_format.space_after = Pt(12)
                skip_mode_icons = True
                continue
                
            # Heading 1
            is_sec_heading = False
            sec_num = None
            m = re.match(r'^(\d+)\.\s+(.*)$', text)
            if m:
                is_sec_heading = True
                sec_num = m.group(1)
                
            if is_sec_heading:
                dest_p = dest_doc.add_paragraph()
                dest_p.paragraph_format.space_before = Pt(18)
                dest_p.paragraph_format.space_after = Pt(6)
                dest_p.paragraph_format.keep_with_next = True
                
                run = dest_p.add_run(text)
                run.font.name = 'Segoe UI'
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 48, 135)
                set_paragraph_left_border(dest_p, color_hex="E1B924", size=32, space=10)
                
                add_bookmark(dest_p, bookmark_counter, f"sec_{sec_num}")
                bookmark_counter += 1
                active_heading_level = None
                continue
                
            # Heading 2 or 3
            is_sub_heading = False
            heading_level = None
            if active_heading_level is not None:
                is_sub_heading = True
                heading_level = 2
                active_heading_level = None
            elif style_name.startswith('Heading'):
                is_sub_heading = True
                try:
                    level = int(style_name.replace('Heading', '').strip())
                    heading_level = 2 if level <= 3 else 3
                except ValueError:
                    heading_level = 2
                    
            if is_sub_heading:
                dest_p = dest_doc.add_paragraph()
                dest_p.paragraph_format.keep_with_next = True
                if heading_level == 2:
                    dest_p.paragraph_format.space_before = Pt(14)
                    dest_p.paragraph_format.space_after = Pt(4)
                    run = dest_p.add_run(text)
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(13)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 48, 135)
                else:
                    dest_p.paragraph_format.space_before = Pt(10)
                    dest_p.paragraph_format.space_after = Pt(2)
                    run = dest_p.add_run(text)
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(11.5)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(43, 43, 43)
                continue
                
            # Lists and Paragraphs
            is_bullet = is_list_paragraph(p)
            if is_bullet:
                dest_p = dest_doc.add_paragraph(style='List Bullet')
                dest_p.paragraph_format.space_before = Pt(0)
                dest_p.paragraph_format.space_after = Pt(3)
                dest_p.paragraph_format.line_spacing = 1.15
                copy_paragraph_runs(p, dest_p, font_name="Segoe UI", font_size=10.5, color_rgb=(43, 43, 43))
            else:
                dest_p = dest_doc.add_paragraph()
                is_note = text.startswith(('*', '💡', '⚠️'))
                if is_note:
                    dest_p.paragraph_format.space_before = Pt(4)
                    dest_p.paragraph_format.space_after = Pt(8)
                    dest_p.paragraph_format.line_spacing = 1.15
                    if text.startswith('💡'):
                        set_paragraph_left_border(dest_p, color_hex="003087", size=16, space=6)
                        copy_paragraph_runs(p, dest_p, font_name="Segoe UI", font_size=10, color_rgb=(50, 50, 50))
                    elif text.startswith('⚠️'):
                        set_paragraph_left_border(dest_p, color_hex="E1B924", size=16, space=6)
                        copy_paragraph_runs(p, dest_p, font_name="Segoe UI", font_size=10, color_rgb=(50, 50, 50))
                    else:
                        copy_paragraph_runs(p, dest_p, font_name="Segoe UI", font_size=9.5, color_rgb=(100, 100, 100))
                else:
                    dest_p.paragraph_format.space_before = Pt(0)
                    dest_p.paragraph_format.space_after = Pt(6)
                    dest_p.paragraph_format.line_spacing = 1.15
                    copy_paragraph_runs(p, dest_p, font_name="Segoe UI", font_size=11, color_rgb=(43, 43, 43))
                    
        elif child.tag.endswith('tbl'):
            if started:
                src_table = Table(child, src_doc)
                num_rows = len(src_table.rows)
                num_cols = len(src_table.columns)
                
                dest_table = dest_doc.add_table(rows=num_rows, cols=num_cols)
                
                for r_idx in range(num_rows):
                    for c_idx in range(num_cols):
                        src_cell = src_table.cell(r_idx, c_idx)
                        dest_cell = dest_table.cell(r_idx, c_idx)
                        dest_cell.text = ""
                        for p_idx, src_p in enumerate(src_cell.paragraphs):
                            if p_idx == 0:
                                dest_p = dest_cell.paragraphs[0]
                            else:
                                dest_p = dest_cell.add_paragraph()
                            copy_paragraph_runs(src_p, dest_p, font_name="Segoe UI", font_size=9.5)
                
                col_alignments = get_table_alignments(table_index, num_cols)
                style_table(dest_table, col_alignments=col_alignments)
                
                widths = get_table_widths(table_index)
                if widths:
                    set_table_widths(dest_table, widths)
                    
                table_index += 1
                p_space = dest_doc.add_paragraph()
                p_space.paragraph_format.space_before = Pt(0)
                p_space.paragraph_format.space_after = Pt(12)

    dest_doc.save(dest_path)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python beautify_docx.py <input.docx> <output.docx>")
    else:
        beautify_document(sys.argv[1], sys.argv[2])
